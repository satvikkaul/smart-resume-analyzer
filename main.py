from fastapi import FastAPI, UploadFile, File, HTTPException, Header
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware

from rapidfuzz import fuzz, process

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

import fitz  # PyMuPDF
import re, io, json, os
from pathlib import Path
import docx
from docx.opc.exceptions import OpcError
from typing import Set, Dict, List

# --------------------------------------------------------------------------------------
# App setup
# --------------------------------------------------------------------------------------
app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")

# Enable CORS for local React dev (adjust as needed)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --------------------------------------------------------------------------------------
# Config / constants
# --------------------------------------------------------------------------------------
MAX_BYTES = 5 * 1024 * 1024   # 5 MB per file
FUZZ_THRESHOLD = 85           # 80–90 typical

CONFIG_PATH = Path("skills_config.json")
ADMIN_TOKEN = os.getenv("ADMIN_TOKEN", "change-me")  # set env var in prod

# Filled by loader at startup
STOPWORDS: Set[str] = set()
SKILL_LEXICON: Set[str] = set()         # base skill tokens
SKILL_ALIASES: Dict[str, List[str]] = {}  # base -> [alias, ...]

# --------------------------------------------------------------------------------------
# Regex (precompiled)
# --------------------------------------------------------------------------------------
EMAIL_RE   = re.compile(r"\S+@\S+")
PHONE_RE   = re.compile(r"\b(?:\+?\d[\d\-\s]{7,}\d)\b")
URL_RE     = re.compile(r"https?://\S+|www\.\S+")
CLEAN_RE   = re.compile(r"[^a-z0-9\+\#\.\-\/\s]")
SPACE_RE   = re.compile(r"\s+")
TOKEN_RE   = re.compile(r"[a-z0-9\+\#\.\-\/]{2,}")

# --------------------------------------------------------------------------------------
# Config loader
# --------------------------------------------------------------------------------------
def load_skills_config(path: Path = CONFIG_PATH):
    global STOPWORDS, SKILL_LEXICON, SKILL_ALIASES
    if not path.exists():
        raise RuntimeError(f"Config not found: {path.resolve()}")
    data = json.loads(path.read_text(encoding="utf-8"))

    STOPWORDS = set((w.lower().strip() for w in data.get("stopwords", [])))

    SKILL_LEXICON = set()
    SKILL_ALIASES = {}
    for item in data.get("skills", []):
        base = item["name"].lower().strip()
        SKILL_LEXICON.add(base)
        alts = [a.lower().strip() for a in item.get("aliases", [])]
        SKILL_ALIASES[base] = alts

    # Optional: include alias tokens in lexicon to simplify filtering
    for _, alts in SKILL_ALIASES.items():
        for a in alts:
            SKILL_LEXICON.add(a)

# Load once at startup
load_skills_config()

# --------------------------------------------------------------------------------------
# Helpers: I/O and extraction
# --------------------------------------------------------------------------------------
async def _read_bytes(f: UploadFile) -> bytes:
    await f.seek(0)
    data = await f.read()
    await f.seek(0)
    if len(data) > MAX_BYTES:
        raise HTTPException(status_code=413, detail="File too large (>5MB)")
    return data

async def extract_text(file: UploadFile) -> str:
    name = (file.filename or "").lower()
    data = await _read_bytes(file)

    # .txt
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")

    # .pdf
    if name.endswith(".pdf"):
        try:
            with fitz.open(stream=data, filetype="pdf") as doc:
                if doc.is_encrypted:
                    try:
                        doc.authenticate("")
                    except Exception:
                        raise HTTPException(status_code=400, detail="Encrypted PDF not supported")
                text_chunks = [page.get_text("text") for page in doc]
        except (fitz.FileDataError, RuntimeError):
            raise HTTPException(status_code=400, detail="Invalid or unreadable PDF")
        return "\n".join(t for t in text_chunks if t and t.strip())

    # .docx
    if name.endswith(".docx"):
        try:
            d = docx.Document(io.BytesIO(data))
        except OpcError:
            raise HTTPException(status_code=400, detail="Invalid DOCX file")

        parts: List[str] = []
        # paragraphs
        for p in d.paragraphs:
            txt = (p.text or "").strip()
            if txt:
                parts.append(txt)
        # tables
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    # unsupported
    raise HTTPException(status_code=415, detail="Unsupported file type. Use .txt, .pdf, or .docx.")

# --------------------------------------------------------------------------------------
# Text processing
# --------------------------------------------------------------------------------------
def normalize(s: str) -> str:
    s = s.lower()
    s = EMAIL_RE.sub(" ", s)
    s = PHONE_RE.sub(" ", s)
    s = URL_RE.sub(" ", s)
    s = CLEAN_RE.sub(" ", s)
    s = SPACE_RE.sub(" ", s).strip()
    return s

def tokenize_words(s: str) -> List[str]:
    return TOKEN_RE.findall(s)

def extract_candidate_keywords(jd_text: str) -> List[str]:
    # JD tokens filtered to skill lexicon or known base skills
    tokens = [t for t in tokenize_words(jd_text) if t not in STOPWORDS]
    tokens = [t for t in tokens if (t in SKILL_LEXICON) or (t in SKILL_ALIASES)]
    freq: Dict[str, int] = {}
    for t in tokens:
        freq[t] = freq.get(t, 0) + 1
    top = sorted(freq.items(), key=lambda x: x[1], reverse=True)[:50]
    return [w for w, _ in top]

# --------------------------------------------------------------------------------------
# Matching and scoring
# --------------------------------------------------------------------------------------
def fuzzy_match_term(term: str, resume_tokens: List[str]) -> bool:
    # exact first
    if term in resume_tokens:
        return True
    # aliases and fuzzy
    candidates = [term] + SKILL_ALIASES.get(term, [])
    for cand in candidates:
        m = process.extractOne(cand, resume_tokens, scorer=fuzz.token_set_ratio)
        if m and m[1] >= FUZZ_THRESHOLD:
            return True
    return False

def _lexicon_vocabulary() -> Dict[str, int]:
    # Build a fixed vocabulary from skills (keep tokens that match TOKEN_RE)
    vocab_terms = {t for t in SKILL_LEXICON if TOKEN_RE.fullmatch(t)}
    # Stable ordering
    return {term: i for i, term in enumerate(sorted(vocab_terms))}

def tfidf_cosine_lexicon(resume_text: str, jd_text: str) -> float:
    vocab = _lexicon_vocabulary()
    if not vocab:
        return 0.0
    try:
        vectorizer = TfidfVectorizer(
            ngram_range=(1, 2),
            lowercase=False,  # already normalized
            token_pattern=r"[a-z0-9\+\#\.\-\/]{2,}",
            stop_words=None,
            vocabulary=vocab,
            sublinear_tf=True,
        )
        X = vectorizer.fit_transform([jd_text, resume_text])
        cos = cosine_similarity(X[0], X[1])[0][0]
        return float(round(cos, 3))
    except ValueError:
        return 0.0

def keyword_match(resume_text: str, jd_text: str) -> Dict:
    jd_keys = extract_candidate_keywords(jd_text)
    resume_tokens_list = list(set(tokenize_words(resume_text)))

    matched, missing = [], []
    for k in jd_keys:
        if fuzzy_match_term(k, resume_tokens_list):
            matched.append(k)
        else:
            missing.append(k)

    coverage = round(len(matched) / max(1, len(jd_keys)), 3)
    cosine_similarity_score = tfidf_cosine_lexicon(resume_text, jd_text)
    overall_score = round(0.75 * cosine_similarity_score + 0.25 * coverage, 3)

    return {
        "overall_score": overall_score,
        "cosine_similarity": cosine_similarity_score,
        "coverage": coverage,
        "matched": matched[:50],
        "missing_sample": missing[:20],
    }

# --------------------------------------------------------------------------------------
# Endpoints
# --------------------------------------------------------------------------------------
@app.post("/analyze")
async def analyze_files(
    resume: UploadFile = File(...),
    job_description: UploadFile = File(...)
):
    resume_text = await extract_text(resume)
    jd_text = await extract_text(job_description)
    return {
        "resume_preview": resume_text[:400],
        "jd_preview": jd_text[:400],
        "resume_chars": len(resume_text),
        "jd_chars": len(jd_text),
    }

@app.post("/score")
async def score(
    resume: UploadFile = File(...),
    job_description: UploadFile = File(...)
):
    resume_raw = await extract_text(resume)
    jd_raw = await extract_text(job_description)
    resume_text = normalize(resume_raw)
    jd_text = normalize(jd_raw)
    return keyword_match(resume_text, jd_text)

@app.post("/admin/reload-skills")
def reload_skills(x_admin_token: str = Header(None)):
    if x_admin_token != ADMIN_TOKEN:
        raise HTTPException(status_code=401, detail="Unauthorized")
    try:
        load_skills_config()
        return {
            "status": "ok",
            "skills_in_lexicon": len(SKILL_LEXICON),
            "bases_with_aliases": len(SKILL_ALIASES),
            "stopwords": len(STOPWORDS),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Reload failed: {e}")
